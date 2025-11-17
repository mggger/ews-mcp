"""Attachment management tools for EWS MCP Server."""

from typing import Any, Dict, List
import base64
import io
from pathlib import Path

from .base import BaseTool
from ..exceptions import ToolExecutionError
from ..utils import format_success_response, safe_get, find_message_across_folders


class ListAttachmentsTool(BaseTool):
    """Tool for listing email attachments."""

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": "list_attachments",
            "description": "List all attachments for a specific email message",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Email message ID"
                    },
                    "include_inline": {
                        "type": "boolean",
                        "description": "Include inline attachments (images embedded in email)",
                        "default": True
                    }
                },
                "required": ["message_id"]
            }
        }

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """List attachments for an email."""
        message_id = kwargs.get("message_id")
        include_inline = kwargs.get("include_inline", True)

        if not message_id:
            raise ToolExecutionError("message_id is required")

        try:
            # Find message across all folders (including custom subfolders)
            message = find_message_across_folders(self.ews_client, message_id)

            # List attachments
            attachments = []
            if hasattr(message, 'attachments') and message.attachments:
                for attachment in message.attachments:
                    # Check if it's inline and whether to include it
                    is_inline = safe_get(attachment, 'is_inline', False)

                    if not include_inline and is_inline:
                        continue

                    # Extract attachment ID properly
                    att_id = safe_get(attachment, 'attachment_id', '')
                    if hasattr(att_id, 'id'):
                        # AttachmentId object
                        att_id = att_id.id
                    elif isinstance(att_id, dict):
                        # Dictionary
                        att_id = att_id.get('id', '')
                    else:
                        # String or other
                        att_id = str(att_id) if att_id else ''

                    attachment_info = {
                        "id": att_id,
                        "name": safe_get(attachment, 'name', 'unnamed'),
                        "size": safe_get(attachment, 'size', 0),
                        "content_type": safe_get(attachment, 'content_type', 'application/octet-stream'),
                        "is_inline": is_inline,
                        "content_id": safe_get(attachment, 'content_id')
                    }
                    attachments.append(attachment_info)

            self.logger.info(f"Found {len(attachments)} attachment(s) for message {message_id}")

            return format_success_response(
                f"Found {len(attachments)} attachment(s)",
                message_id=message_id,
                attachments=attachments,
                count=len(attachments)
            )

        except ToolExecutionError:
            raise
        except Exception as e:
            self.logger.error(f"Failed to list attachments: {e}")
            raise ToolExecutionError(f"Failed to list attachments: {e}")


class DownloadAttachmentTool(BaseTool):
    """Tool for downloading email attachments."""

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": "download_attachment",
            "description": "Download an email attachment by message ID and attachment ID",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Email message ID"
                    },
                    "attachment_id": {
                        "type": "string",
                        "description": "Attachment ID (from list_attachments)"
                    },
                    "return_as": {
                        "type": "string",
                        "enum": ["base64", "file_path"],
                        "description": "Return format: base64 encoded string or file path",
                        "default": "base64"
                    },
                    "save_path": {
                        "type": "string",
                        "description": "File path to save attachment (required if return_as is 'file_path')"
                    }
                },
                "required": ["message_id", "attachment_id"]
            }
        }

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """Download an attachment."""
        message_id = kwargs.get("message_id")
        attachment_id = kwargs.get("attachment_id")
        return_as = kwargs.get("return_as", "base64")
        save_path = kwargs.get("save_path")

        if not message_id:
            raise ToolExecutionError("message_id is required")

        if not attachment_id:
            raise ToolExecutionError("attachment_id is required")

        if return_as == "file_path" and not save_path:
            raise ToolExecutionError("save_path is required when return_as is 'file_path'")

        try:
            # Find message across all folders (including custom subfolders)
            message = find_message_across_folders(self.ews_client, message_id)

            # Find the attachment
            attachment = None
            if hasattr(message, 'attachments') and message.attachments:
                for att in message.attachments:
                    att_id = safe_get(att, 'attachment_id', {})
                    if isinstance(att_id, dict):
                        att_id = att_id.get('id', '')

                    if str(att_id) == str(attachment_id):
                        attachment = att
                        break

            if not attachment:
                raise ToolExecutionError(f"Attachment not found: {attachment_id}")

            # Get attachment content
            content = safe_get(attachment, 'content', b'')
            if not content:
                raise ToolExecutionError("Attachment content is empty")

            attachment_name = safe_get(attachment, 'name', 'attachment')

            # Return based on format
            if return_as == "base64":
                # Encode content as base64
                content_b64 = base64.b64encode(content).decode('utf-8')

                self.logger.info(f"Downloaded attachment {attachment_name} ({len(content)} bytes)")

                return format_success_response(
                    "Attachment downloaded successfully",
                    message_id=message_id,
                    attachment_id=attachment_id,
                    name=attachment_name,
                    size=len(content),
                    content_type=safe_get(attachment, 'content_type', 'application/octet-stream'),
                    content_base64=content_b64
                )

            elif return_as == "file_path":
                # Save to file
                file_path = Path(save_path)

                # Create parent directories if they don't exist
                file_path.parent.mkdir(parents=True, exist_ok=True)

                # Write content to file
                with open(file_path, 'wb') as f:
                    f.write(content)

                self.logger.info(f"Saved attachment {attachment_name} to {file_path}")

                return format_success_response(
                    "Attachment saved successfully",
                    message_id=message_id,
                    attachment_id=attachment_id,
                    name=attachment_name,
                    size=len(content),
                    content_type=safe_get(attachment, 'content_type', 'application/octet-stream'),
                    file_path=str(file_path)
                )

        except ToolExecutionError:
            raise
        except Exception as e:
            self.logger.error(f"Failed to download attachment: {e}")
            raise ToolExecutionError(f"Failed to download attachment: {e}")


class AddAttachmentTool(BaseTool):
    """Tool for adding attachments to draft or existing emails."""

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": "add_attachment",
            "description": "Add an attachment to a draft or existing email message",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Email message ID to add attachment to"
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Local file path to attach"
                    },
                    "file_content": {
                        "type": "string",
                        "description": "Base64 encoded file content (alternative to file_path)"
                    },
                    "file_name": {
                        "type": "string",
                        "description": "File name for the attachment (required if using file_content)"
                    },
                    "content_type": {
                        "type": "string",
                        "description": "MIME content type (e.g., 'application/pdf')",
                        "default": "application/octet-stream"
                    },
                    "is_inline": {
                        "type": "boolean",
                        "description": "Whether attachment is inline (embedded in body)",
                        "default": False
                    }
                },
                "required": ["message_id"]
            }
        }

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """Add attachment to email."""
        message_id = kwargs.get("message_id")
        file_path = kwargs.get("file_path")
        file_content_b64 = kwargs.get("file_content")
        file_name = kwargs.get("file_name")
        content_type = kwargs.get("content_type", "application/octet-stream")
        is_inline = kwargs.get("is_inline", False)

        if not message_id:
            raise ToolExecutionError("message_id is required")

        if not file_path and not file_content_b64:
            raise ToolExecutionError("Either file_path or file_content is required")

        if file_content_b64 and not file_name:
            raise ToolExecutionError("file_name is required when using file_content")

        try:
            from exchangelib import FileAttachment
            import base64
            from pathlib import Path

            # Find the message in Drafts folder (most common use case)
            message = None
            folders_to_search = [
                self.ews_client.account.drafts,
                self.ews_client.account.inbox,
                self.ews_client.account.sent
            ]

            for folder in folders_to_search:
                try:
                    message = folder.get(id=message_id)
                    if message:
                        break
                except Exception:
                    continue

            if not message:
                raise ToolExecutionError(f"Message not found: {message_id}")

            # Get file content and name
            if file_path:
                path = Path(file_path)
                if not path.exists():
                    raise ToolExecutionError(f"File not found: {file_path}")

                with open(path, 'rb') as f:
                    file_content = f.read()

                if not file_name:
                    file_name = path.name
            else:
                # Decode base64 content
                file_content = base64.b64decode(file_content_b64)

            # Create attachment
            attachment = FileAttachment(
                name=file_name,
                content=file_content,
                content_type=content_type,
                is_inline=is_inline
            )

            # Add attachment to message
            if not hasattr(message, 'attachments'):
                message.attachments = []

            message.attachments.append(attachment)
            message.save()

            self.logger.info(f"Added attachment '{file_name}' to message {message_id}")

            return format_success_response(
                f"Attachment '{file_name}' added successfully",
                message_id=message_id,
                attachment_name=file_name,
                attachment_size=len(file_content),
                content_type=content_type,
                is_inline=is_inline
            )

        except ToolExecutionError:
            raise
        except Exception as e:
            self.logger.error(f"Failed to add attachment: {e}")
            raise ToolExecutionError(f"Failed to add attachment: {e}")


class DeleteAttachmentTool(BaseTool):
    """Tool for removing attachments from email messages."""

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": "delete_attachment",
            "description": "Remove an attachment from an email message",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Email message ID containing the attachment"
                    },
                    "attachment_id": {
                        "type": "string",
                        "description": "Attachment ID to remove"
                    },
                    "attachment_name": {
                        "type": "string",
                        "description": "Attachment name to remove (alternative to attachment_id)"
                    }
                },
                "required": ["message_id"]
            }
        }

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """Delete attachment from email."""
        message_id = kwargs.get("message_id")
        attachment_id = kwargs.get("attachment_id")
        attachment_name = kwargs.get("attachment_name")

        if not message_id:
            raise ToolExecutionError("message_id is required")

        if not attachment_id and not attachment_name:
            raise ToolExecutionError("Either attachment_id or attachment_name is required")

        try:
            # Find the message
            message = None
            folders_to_search = [
                self.ews_client.account.drafts,
                self.ews_client.account.inbox,
                self.ews_client.account.sent
            ]

            for folder in folders_to_search:
                try:
                    message = folder.get(id=message_id)
                    if message:
                        break
                except Exception:
                    continue

            if not message:
                raise ToolExecutionError(f"Message not found: {message_id}")

            # Find the attachment to delete
            if not hasattr(message, 'attachments') or not message.attachments:
                raise ToolExecutionError("Message has no attachments")

            attachment_to_delete = None
            deleted_name = None

            for att in message.attachments:
                # Check by ID
                if attachment_id:
                    att_id = safe_get(att, 'attachment_id', {})
                    if isinstance(att_id, dict):
                        att_id = att_id.get('id', '')

                    if str(att_id) == str(attachment_id):
                        attachment_to_delete = att
                        deleted_name = safe_get(att, 'name', 'Unknown')
                        break

                # Check by name
                elif attachment_name:
                    att_name = safe_get(att, 'name', '')
                    if att_name == attachment_name:
                        attachment_to_delete = att
                        deleted_name = att_name
                        break

            if not attachment_to_delete:
                raise ToolExecutionError(f"Attachment not found")

            # Remove the attachment
            message.attachments.remove(attachment_to_delete)
            message.save()

            self.logger.info(f"Deleted attachment '{deleted_name}' from message {message_id}")

            return format_success_response(
                f"Attachment '{deleted_name}' deleted successfully",
                message_id=message_id,
                attachment_name=deleted_name
            )

        except ToolExecutionError:
            raise
        except Exception as e:
            self.logger.error(f"Failed to delete attachment: {e}")
            raise ToolExecutionError(f"Failed to delete attachment: {e}")


class ReadAttachmentTool(BaseTool):
    """Tool for reading and extracting text content from email attachments."""

    def get_schema(self) -> Dict[str, Any]:
        return {
            "name": "read_attachment",
            "description": "Extract text content from email attachments (PDF, DOCX, XLSX, TXT). Supports Arabic (UTF-8) text.",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Email message ID"
                    },
                    "attachment_name": {
                        "type": "string",
                        "description": "Name of the attachment to read"
                    },
                    "extract_tables": {
                        "type": "boolean",
                        "description": "Extract tables from documents (PDF, DOCX, XLSX)",
                        "default": False
                    },
                    "max_pages": {
                        "type": "integer",
                        "description": "Maximum number of pages to extract (for PDFs)",
                        "default": 50
                    }
                },
                "required": ["message_id", "attachment_name"]
            }
        }

    async def execute(self, **kwargs) -> Dict[str, Any]:
        """Read and extract text from attachment."""
        message_id = kwargs.get("message_id")
        attachment_name = kwargs.get("attachment_name")
        extract_tables = kwargs.get("extract_tables", False)
        max_pages = kwargs.get("max_pages", 50)

        if not message_id:
            raise ToolExecutionError("message_id is required")

        if not attachment_name:
            raise ToolExecutionError("attachment_name is required")

        try:
            # Find message across all folders (including custom subfolders)
            message = find_message_across_folders(self.ews_client, message_id)

            # Find the attachment by name
            attachment = None
            if hasattr(message, 'attachments') and message.attachments:
                for att in message.attachments:
                    if safe_get(att, 'name', '') == attachment_name:
                        attachment = att
                        break

            if not attachment:
                raise ToolExecutionError(f"Attachment '{attachment_name}' not found")

            # Get attachment content
            content = safe_get(attachment, 'content', b'')
            if not content:
                raise ToolExecutionError("Attachment content is empty")

            # Determine file type from extension
            file_ext = attachment_name.lower().split('.')[-1] if '.' in attachment_name else ''

            # Extract text based on file type
            extracted_text = ""
            file_type = file_ext

            if file_ext == 'pdf':
                extracted_text = self._read_pdf(content, extract_tables, max_pages)
            elif file_ext == 'docx':
                extracted_text = self._read_docx(content, extract_tables)
            elif file_ext in ['xlsx', 'xls']:
                extracted_text = self._read_excel(content)
            elif file_ext == 'txt':
                extracted_text = content.decode('utf-8', errors='replace')
            else:
                raise ToolExecutionError(
                    f"Unsupported file type: {file_ext}. Supported: PDF, DOCX, XLSX, TXT"
                )

            self.logger.info(
                f"Extracted {len(extracted_text)} characters from {attachment_name}"
            )

            return format_success_response(
                "Text extracted successfully",
                message_id=message_id,
                file_name=attachment_name,
                file_type=file_type,
                file_size=len(content),
                content_length=len(extracted_text),
                content=extracted_text,
                supports_arabic=True  # UTF-8 encoding supports Arabic
            )

        except ToolExecutionError:
            raise
        except Exception as e:
            self.logger.error(f"Failed to read attachment: {e}")
            raise ToolExecutionError(f"Failed to read attachment: {e}")

    def _read_pdf(self, content: bytes, extract_tables: bool, max_pages: int) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            raise ToolExecutionError(
                "pdfplumber not installed. Run: pip install pdfplumber>=0.10.0"
            )

        text_parts = []
        pdf_bytes = io.BytesIO(content)

        try:
            with pdfplumber.open(pdf_bytes) as pdf:
                total_pages = len(pdf.pages)
                pages_to_process = min(total_pages, max_pages)

                for page_num in range(pages_to_process):
                    page = pdf.pages[page_num]
                    text_parts.append(f"--- Page {page_num + 1} of {total_pages} ---")

                    # Extract text (supports UTF-8/Arabic)
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # Extract tables if requested
                    if extract_tables:
                        tables = page.extract_tables()
                        for table_num, table in enumerate(tables, 1):
                            text_parts.append(f"\n[Table {table_num}]")
                            for row in table:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)

                if total_pages > max_pages:
                    text_parts.append(f"\n... (truncated at {max_pages} pages, {total_pages - max_pages} pages omitted)")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise ToolExecutionError(f"Failed to read PDF: {str(e)}")

    def _read_docx(self, content: bytes, extract_tables: bool) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ToolExecutionError(
                "python-docx not installed. Run: pip install python-docx>=1.0.0"
            )

        text_parts = []
        docx_bytes = io.BytesIO(content)

        try:
            doc = Document(docx_bytes)

            # Extract paragraphs (supports UTF-8/Arabic)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables if requested
            if extract_tables:
                for table_num, table in enumerate(doc.tables, 1):
                    text_parts.append(f"\n[Table {table_num}]")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            raise ToolExecutionError(f"Failed to read DOCX: {str(e)}")

    def _read_excel(self, content: bytes) -> str:
        """Extract data from Excel using openpyxl."""
        try:
            import openpyxl
        except ImportError:
            raise ToolExecutionError(
                "openpyxl not installed. Run: pip install openpyxl>=3.1.0"
            )

        text_parts = []
        excel_bytes = io.BytesIO(content)

        try:
            # Load workbook (data_only=True to get calculated values)
            workbook = openpyxl.load_workbook(excel_bytes, data_only=True)

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")

                # Extract all rows
                for row in sheet.iter_rows(values_only=True):
                    # Skip empty rows
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(
                            str(cell) if cell is not None else ""
                            for cell in row
                        )
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line between sheets

            return "\n".join(text_parts)

        except Exception as e:
            raise ToolExecutionError(f"Failed to read Excel: {str(e)}")
